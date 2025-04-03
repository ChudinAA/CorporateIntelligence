import os
import logging
import tempfile
import PyPDF2
import docx
import pandas as pd
import re
from flask import current_app
from werkzeug.utils import secure_filename
from models import Document, DocumentChunk
from app import db
from vector_store import VectorStore
from llm_integration import LLMService

class DocumentProcessor:
    """Handles document processing, text extraction, and chunking."""
    
    def __init__(self):
        self.vector_store = VectorStore()
        self.llm_service = LLMService()
        self.logger = logging.getLogger(__name__)
        
        # Create upload directory if it doesn't exist
        os.makedirs(current_app.config['UPLOAD_FOLDER'], exist_ok=True)
    
    def process_uploaded_file(self, file, user_id):
        """
        Process an uploaded file and store in vector database.
        
        Args:
            file: File object from Flask request
            user_id: ID of the user uploading the file
            
        Returns:
            dict: Processing status information
        """
        try:
            # Validate file
            if not self._is_allowed_file(file.filename):
                return {"success": False, "error": "File type not supported"}
            
            # Secure filename and save file
            original_filename = file.filename
            secure_name = secure_filename(original_filename)
            filename = f"{user_id}_{secure_name}"
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Get file type
            file_type = original_filename.rsplit('.', 1)[1].lower()
            
            # Create document record
            document = Document(
                filename=filename,
                original_filename=original_filename,
                file_type=file_type,
                file_size=file_size,
                user_id=user_id,
                processed=False
            )
            db.session.add(document)
            db.session.commit()
            
            # Extract text from file
            text = self._extract_text(file_path, file_type)
            
            # If text extraction failed
            if not text:
                document.processing_error = "Failed to extract text from document"
                db.session.commit()
                return {
                    "success": False, 
                    "error": "Failed to extract text from document",
                    "document_id": document.id
                }
            
            # Create chunks from text
            chunks = self._create_chunks(text, chunk_size=1000, chunk_overlap=100)
            
            # Store chunks in database and vector store
            chunk_objects = []
            for i, chunk_text in enumerate(chunks):
                # Create embedding for chunk
                embedding = self.llm_service.get_embedding(chunk_text)
                
                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_text=chunk_text,
                    chunk_index=i
                )
                db.session.add(chunk)
                db.session.flush()  # Get ID without committing
                
                # Add to vector store
                vector_id = self.vector_store.add_embedding(
                    embedding=embedding,
                    doc_chunk_id=chunk.id,
                    document_id=document.id,
                    user_id=user_id,
                    metadata={
                        "file_type": file_type,
                        "chunk_index": i,
                        "original_filename": original_filename
                    }
                )
                
                # Update chunk with vector ID
                chunk.vector_id = vector_id
                chunk_objects.append(chunk)
            
            # Mark document as processed
            document.processed = True
            db.session.commit()
            
            return {
                "success": True,
                "document_id": document.id,
                "chunks_count": len(chunks),
                "document_name": original_filename
            }
            
        except Exception as e:
            db.session.rollback()
            self.logger.error(f"Document processing error: {str(e)}", exc_info=True)
            
            # Update document with error if it was created
            try:
                if 'document' in locals() and document.id:
                    document.processing_error = str(e)[:255]  # Truncate if too long
                    db.session.commit()
            except:
                pass
                
            return {"success": False, "error": str(e)}
    
    def _is_allowed_file(self, filename):
        """Check if file type is allowed."""
        allowed_extensions = current_app.config['ALLOWED_EXTENSIONS']
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
    
    def _extract_text(self, file_path, file_type):
        """Extract text from different file formats."""
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type == 'txt':
                return self._extract_from_txt(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            elif file_type in ['xlsx', 'csv']:
                return self._extract_from_spreadsheet(file_path, file_type)
            else:
                self.logger.error(f"Unsupported file type for extraction: {file_type}")
                return None
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_type} file: {str(e)}", exc_info=True)
            return None
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() + "\n"
        return text
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def _extract_from_spreadsheet(self, file_path, file_type):
        """Extract text from Excel/CSV file."""
        if file_type == 'xlsx':
            df = pd.read_excel(file_path)
        else:  # CSV
            df = pd.read_csv(file_path)
        
        # Convert DataFrame to string representation
        text = df.to_string(index=False)
        return text
    
    def _create_chunks(self, text, chunk_size=1000, chunk_overlap=100):
        """
        Split text into overlapping chunks.
        
        Args:
            text (str): The text to split
            chunk_size (int): Maximum size of each chunk
            chunk_overlap (int): Amount of overlap between chunks
            
        Returns:
            list: List of text chunks
        """
        if not text:
            return []
        
        # Split text into sentences (rough heuristic)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If adding this sentence exceeds chunk size and we have content, create a new chunk
            if current_size + sentence_size > chunk_size and current_chunk:
                # Join the current chunk and add to chunks list
                chunks.append(" ".join(current_chunk))
                
                # Keep some sentences for overlap
                overlap_sentences = []
                overlap_size = 0
                
                # Work backwards to create overlap
                for s in reversed(current_chunk):
                    if overlap_size + len(s) <= chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s) + 1  # +1 for space
                    else:
                        break
                
                # Start new chunk with overlap
                current_chunk = overlap_sentences
                current_size = overlap_size
            
            # Add the current sentence
            current_chunk.append(sentence)
            current_size += sentence_size + 1  # +1 for space
        
        # Add the last chunk if it has content
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
